"""
Universal File Ingestion & Parsing Engine for ARGUS AI Platform.
Supports CSV, Excel, JSON, XML, TXT, PDF, Word (Docx), PowerPoint (Pptx), ZIP, and Images.
"""
import io
import json
import zipfile
import xml.etree.ElementTree as ET
import pandas as pd
from typing import Dict, Any, Tuple, List, Optional
from pathlib import Path
from PIL import Image
import pypdf
import docx
from core.logger import logger
from core.exceptions import FileProcessingException
from utils.security import sanitize_filename


class FileIngestionService:
    """Universal Engine to inspect, parse, and convert files into DataFrames or structured text."""

    SUPPORTED_EXTENSIONS = {
        "csv": ["csv", "tsv", "txt"],
        "excel": ["xlsx", "xls", "xlsm"],
        "json": ["json", "jsonl"],
        "xml": ["xml"],
        "pdf": ["pdf"],
        "word": ["docx", "doc"],
        "pptx": ["pptx"],
        "archive": ["zip"],
        "image": ["png", "jpg", "jpeg", "bmp", "tiff", "webp"]
    }

    @classmethod
    def process_uploaded_file(cls, uploaded_file) -> Tuple[Optional[pd.DataFrame], Dict[str, Any]]:
        """
        Main entry point for handling raw uploaded file objects from Streamlit.
        Returns (pd.DataFrame or None, metadata_dict).
        """
        filename = sanitize_filename(uploaded_file.name)
        ext = filename.split(".")[-1].lower() if "." in filename else ""
        file_bytes = uploaded_file.getvalue()
        file_size = len(file_bytes)

        logger.info(f"Ingesting file '{filename}' ({file_size} bytes, type '{ext}')")

        metadata: Dict[str, Any] = {
            "filename": filename,
            "extension": ext,
            "size_bytes": file_size,
            "parsed_type": "unknown",
            "text_content": None,
            "error": None
        }

        try:
            # 1. CSV / TSV / Delimited Text
            if ext in cls.SUPPORTED_EXTENSIONS["csv"]:
                df = cls._parse_csv(file_bytes)
                metadata["parsed_type"] = "dataframe"
                return df, metadata

            # 2. Excel Spreadsheets
            elif ext in cls.SUPPORTED_EXTENSIONS["excel"]:
                df = cls._parse_excel(file_bytes)
                metadata["parsed_type"] = "dataframe"
                return df, metadata

            # 3. JSON Data
            elif ext in cls.SUPPORTED_EXTENSIONS["json"]:
                df = cls._parse_json(file_bytes)
                metadata["parsed_type"] = "dataframe"
                return df, metadata

            # 4. XML Data
            elif ext in cls.SUPPORTED_EXTENSIONS["xml"]:
                df = cls._parse_xml(file_bytes)
                metadata["parsed_type"] = "dataframe"
                return df, metadata

            # 5. PDF Document
            elif ext in cls.SUPPORTED_EXTENSIONS["pdf"]:
                text_content, df_tables = cls._parse_pdf(file_bytes)
                metadata["parsed_type"] = "document"
                metadata["text_content"] = text_content
                return df_tables, metadata

            # 6. Word Document (Docx)
            elif ext in cls.SUPPORTED_EXTENSIONS["word"]:
                text_content, df = cls._parse_docx(file_bytes)
                metadata["parsed_type"] = "document"
                metadata["text_content"] = text_content
                return df, metadata

            # 7. ZIP Archive
            elif ext in cls.SUPPORTED_EXTENSIONS["archive"]:
                extracted_dfs, zip_meta = cls._parse_zip(file_bytes)
                metadata.update(zip_meta)
                metadata["parsed_type"] = "archive"
                # Combine first available DataFrame
                combined_df = extracted_dfs[0] if extracted_dfs else None
                return combined_df, metadata

            # 8. Images
            elif ext in cls.SUPPORTED_EXTENSIONS["image"]:
                img = Image.open(io.BytesIO(file_bytes))
                metadata["parsed_type"] = "image"
                metadata["image_size"] = img.size
                metadata["image_mode"] = img.mode
                return None, metadata

            else:
                # Default text fallback attempt
                df = cls._parse_csv(file_bytes)
                metadata["parsed_type"] = "dataframe"
                return df, metadata

        except Exception as e:
            logger.error(f"Error processing file '{filename}': {str(e)}")
            metadata["error"] = str(e)
            raise FileProcessingException(f"Failed to process '{filename}': {str(e)}")

    @classmethod
    def _parse_csv(cls, file_bytes: bytes) -> pd.DataFrame:
        """Robust CSV parser with encoding auto-detection fallback."""
        for encoding in ["utf-8", "latin-1", "cp1252", "iso-8859-1"]:
            try:
                # Test comma vs tab vs semicolon
                sample = file_bytes[:4096].decode(encoding, errors="ignore")
                sep = ","
                if sample.count("\t") > sample.count(","):
                    sep = "\t"
                elif sample.count(";") > sample.count(","):
                    sep = ";"
                
                df = pd.read_csv(io.BytesIO(file_bytes), encoding=encoding, sep=sep)
                return df
            except Exception:
                continue
        raise FileProcessingException("Could not parse CSV file with standard encodings.")

    @classmethod
    def _parse_excel(cls, file_bytes: bytes) -> pd.DataFrame:
        """Excel parser extracting primary sheet into DataFrame."""
        try:
            xls = pd.ExcelFile(io.BytesIO(file_bytes))
            sheet_name = xls.sheet_names[0]
            df = pd.read_excel(xls, sheet_name=sheet_name)
            return df
        except Exception as e:
            raise FileProcessingException(f"Excel Parse Error: {str(e)}")

    @classmethod
    def _parse_json(cls, file_bytes: bytes) -> pd.DataFrame:
        """JSON parser supporting list of dicts or nested keys."""
        try:
            data = json.loads(file_bytes.decode("utf-8", errors="ignore"))
            if isinstance(data, list):
                return pd.DataFrame(data)
            elif isinstance(data, dict):
                # Normalize json
                return pd.json_normalize(data)
            else:
                return pd.DataFrame([data])
        except Exception as e:
            raise FileProcessingException(f"JSON Parse Error: {str(e)}")

    @classmethod
    def _parse_xml(cls, file_bytes: bytes) -> pd.DataFrame:
        """XML parser extracting elements into tabular DataFrame."""
        try:
            tree = ET.fromstring(file_bytes.decode("utf-8", errors="ignore"))
            rows = []
            for child in tree:
                row = {}
                for elem in child:
                    row[elem.tag] = elem.text
                if row:
                    rows.append(row)
            if not rows:
                # Flat parse
                rows = [{elem.tag: elem.text for elem in tree}]
            return pd.DataFrame(rows)
        except Exception as e:
            raise FileProcessingException(f"XML Parse Error: {str(e)}")

    @classmethod
    def _parse_pdf(cls, file_bytes: bytes) -> Tuple[str, Optional[pd.DataFrame]]:
        """PDF Parser extracting text content using PyPDF."""
        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            text_lines = []
            for page_idx, page in enumerate(reader.pages):
                extracted = page.extract_text()
                if extracted:
                    text_lines.append(f"--- Page {page_idx + 1} ---\n" + extracted)
            
            full_text = "\n\n".join(text_lines)
            
            # Simple text line DataFrame fallback
            lines = [line.strip() for line in full_text.split("\n") if line.strip()]
            df = pd.DataFrame({"document_line": lines}) if lines else None
            return full_text, df
        except Exception as e:
            raise FileProcessingException(f"PDF Parse Error: {str(e)}")

    @classmethod
    def _parse_docx(cls, file_bytes: bytes) -> Tuple[str, Optional[pd.DataFrame]]:
        """Word Document (.docx) parser."""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text])
            
            # Extract tables if present
            table_rows = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_rows.append(row_data)
            
            if table_rows and len(table_rows) > 1:
                df = pd.DataFrame(table_rows[1:], columns=table_rows[0])
            else:
                lines = [line.strip() for line in full_text.split("\n") if line.strip()]
                df = pd.DataFrame({"paragraph": lines}) if lines else None
                
            return full_text, df
        except Exception as e:
            raise FileProcessingException(f"Word Document Parse Error: {str(e)}")

    @classmethod
    def _parse_zip(cls, file_bytes: bytes) -> Tuple[List[pd.DataFrame], Dict[str, Any]]:
        """ZIP Archive extractor parsing embedded CSV/Excel files."""
        extracted_dfs = []
        file_list = []
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                for filename in z.namelist():
                    if filename.startswith("__MACOSX") or filename.endswith("/"):
                        continue
                    file_list.append(filename)
                    ext = filename.split(".")[-1].lower()
                    if ext in ["csv", "xlsx", "json"]:
                        sub_bytes = z.read(filename)
                        if ext == "csv":
                            extracted_dfs.append(cls._parse_csv(sub_bytes))
                        elif ext == "xlsx":
                            extracted_dfs.append(cls._parse_excel(sub_bytes))
                        elif ext == "json":
                            extracted_dfs.append(cls._parse_json(sub_bytes))
            return extracted_dfs, {"contained_files": file_list, "extracted_count": len(extracted_dfs)}
        except Exception as e:
            raise FileProcessingException(f"ZIP Archive Extract Error: {str(e)}")
